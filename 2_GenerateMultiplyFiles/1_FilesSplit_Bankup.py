# -*- coding: utf-8 -*-
# @Author  :  connor
# @Time    :  2023/9/6 13:44
# @File    :  1_FilesSplit_Bankup.py
# @IDE     :  PyCharm

"""
"""
import docx
import xlrd
import pandas as pd
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


if __name__ == '__main__':

    file = "./待调单-总表.xls"
    pd.set_option('display.float_format', lambda x: '%.2f' % x)
    df = pd.read_excel(file, dtype={"被查账/卡号": str})
    workbook = xlrd.open_workbook(file)  # 文件路径
    worksheet = workbook.sheet_by_index(0)
    nrows = worksheet.nrows  # 获取该表总行数

    flag = 1
    all_data = {}
    all_df = {}
    data = []
    df_data = []
    for i in range(1, nrows):  # 循环打印每一行
        if worksheet.row_values(i)[2] not in all_data:
            all_data[worksheet.row_values(i)[2]] = []
            all_df[worksheet.row_values(i)[2]] = []
        all_data[worksheet.row_values(i)[2]].append(worksheet.row_values(i)[1])
        all_df[worksheet.row_values(i)[2]].append(df[i - 1:i])

    for ik, iv in all_data.items():
        if len(iv) > 50:
            fileNum = int(len(iv) / 50) + 1
            row = 50
            rows = 50
            countNum = 0
            for i in range(fileNum):
                # 创建内存中的word文档对象
                file = docx.Document()

                # 写入若干段落
                # file.add_paragraph("，".join(list(set(iv))))
                paragraph_title = file.add_paragraph()
                run = paragraph_title.add_run(f"附件：{ik}{i + 1}-{len(list(set(iv[countNum: row])))}张")
                run.font.size = Pt(16)
                font = run.font
                font.name = '黑体'
                font._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')

                contents = list(set(iv[countNum: row]))
                paragraph_content = file.add_paragraph()
                paragraph_content_ = file.add_paragraph()
                run_ = paragraph_content.add_run("，".join(contents))
                font = run_.font
                font.size = Pt(14)
                font.name = '华文宋体'
                font._element.rPr.rFonts.set(qn('w:eastAsia'), u'华文宋体')

                run_2 = paragraph_content_.add_run('\n平顶山市公安局卫东分局\n二〇二三年七月二十七日')
                paragraph_content_.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run_2.font.size = Pt(14)
                font = run_2.font
                font.name = '华文宋体'
                font._element.rPr.rFonts.set(qn('w:eastAsia'), u'华文宋体')

                # 保存
                file.save(f"文书附件-{ik}{i + 1}-{len(list(set(iv[countNum: row])))}张.docx")
                countNum += rows
                row += rows
        else:
            # 创建内存中的word文档对象
            file = docx.Document()

            # 写入若干段落
            # file.add_paragraph("，".join(list(set(iv))))
            paragraph_title = file.add_paragraph()
            run = paragraph_title.add_run(f"附件：{ik}-{len(list(set(iv)))}张")
            run.font.size = Pt(16)
            font = run.font
            font.name = '黑体'
            font._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')

            contents = list(set(iv))
            paragraph_content = file.add_paragraph()
            paragraph_content_ = file.add_paragraph()
            run_ = paragraph_content.add_run("，".join(contents))
            font = run_.font
            font.size = Pt(14)
            font.name = '华文宋体'
            font._element.rPr.rFonts.set(qn('w:eastAsia'), u'华文宋体')

            run_2 = paragraph_content_.add_run('\n平顶山市公安局卫东分局\n二〇二三年七月二十七日')
            paragraph_content_.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run_2.font.size = Pt(14)
            font = run_2.font
            font.name = '华文宋体'
            font._element.rPr.rFonts.set(qn('w:eastAsia'), u'华文宋体')

            # 保存
            file.save(f"文书附件-{ik}-{len(list(set(iv)))}张.docx")

    for ik, iv in all_df.items():
        if len(iv) > 50:
            fileNum = int(len(iv) / 50) + 1
            row = 50
            rows = 50
            countNum = 0
            for i in range(fileNum):
                pd.concat(iv[countNum: row]).to_excel(f'文书-{ik}{i + 1}.xls', index=False, engine='openpyxl')
                countNum += rows
                row += rows
        else:
            pd.concat(iv).to_excel(f'文书-{ik}.xls', index=False, engine='openpyxl')