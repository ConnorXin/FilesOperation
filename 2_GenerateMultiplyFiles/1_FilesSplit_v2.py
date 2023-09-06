# -*- coding: utf-8 -*-
# @Author  :  connor
# @Time    :  2023/9/6 16:57
# @File    :  1_FilesSplit_v2.py
# @IDE     :  PyCharm

"""
"""
import os
import re
from operator import itemgetter

import docx
import xlrd
import pandas as pd
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docxtpl import DocxTemplate


if __name__ == '__main__':

    file = "./待调单-总表.xls"  # 总表
    pd.set_option('display.float_format', lambda x: '%.2f' % x)
    df = pd.read_excel(file, dtype={"被查账/卡号": str})
    workbook = xlrd.open_workbook(file)  # 文件路径
    worksheet = workbook.sheet_by_index(0)
    nrows = worksheet.nrows  # 获取该表总行数

    flag = 1
    all_data = {}
    all_df = {}
    data = []
    df_data = []
    no100_df = []
    for i in range(1, nrows):  # 循环打印每一行
        if worksheet.row_values(i)[2] not in all_data:
            all_data[worksheet.row_values(i)[2]] = []
            all_df[worksheet.row_values(i)[2]] = []
        all_data[worksheet.row_values(i)[2]].append(worksheet.row_values(i)[1])
        all_df[worksheet.row_values(i)[2]].append(df[i - 1:i])

    folder_num = 0
    for ik, iv in all_data.items():
        if len(iv) > 100:
            fileNum = int(len(iv) / 100) + 1
            row = 100
            rows = 100
            countNum = 0
            for i in range(fileNum):
                # 创建内存中的word文档对象
                file = docx.Document()
                # 写入若干段落
                # file.add_paragraph("，".join(list(set(iv))))
                paragraph_title = file.add_paragraph()
                run = paragraph_title.add_run(f"附件：{ik}{i + 1}-{len(list(set(iv[countNum: row])))}张")
                run.font.size = Pt(16)
                font = run.font
                font.name = '黑体'
                font._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')

                contents = list(set(iv[countNum: row]))
                paragraph_content = file.add_paragraph()
                paragraph_content_ = file.add_paragraph()
                run_ = paragraph_content.add_run("，".join(contents))
                font = run_.font
                font.size = Pt(14)
                font.name = '华文宋体'
                font._element.rPr.rFonts.set(qn('w:eastAsia'), u'华文宋体')

                run_2 = paragraph_content_.add_run('\n平顶山市公安局卫东分局\n二〇二三年七月二十七日')
                paragraph_content_.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run_2.font.size = Pt(14)
                font = run_2.font
                font.name = '华文宋体'
                font._element.rPr.rFonts.set(qn('w:eastAsia'), u'华文宋体')

                # 保存
                if len(iv[countNum: row]) == 100:
                    os.mkdir(f'./{folder_num + 1}')
                    file.save(f"./{folder_num + 1}/文书附件-{ik}{i + 1}-{len(list(set(iv[countNum: row])))}张.docx")
                    countNum += rows
                    row += rows
                    folder_num += 1
                elif len(iv[countNum: row]) < 100 and len(iv[countNum: row]) > 90:
                    os.mkdir(f'./{folder_num + 1}')
                    file.save(f"./{folder_num + 1}/文书附件-{ik}{i + 1}-{len(list(set(iv[countNum: row])))}张.docx")
                    countNum += rows
                    row += rows
                    folder_num += 1
                else:
                    no100_df.append({'content': file, 'filename': '文书附件-' + str(ik) + str((i + 1)) + '-' + str(len(list(set(iv[countNum: row])))) + '张.docx',
                                     'len': len(list(set(iv[countNum: row])))})
        else:
            # 创建内存中的word文档对象
            file = docx.Document()

            # 写入若干段落
            # file.add_paragraph("，".join(list(set(iv))))
            paragraph_title = file.add_paragraph()
            run = paragraph_title.add_run(f"附件：{ik}-{len(list(set(iv)))}张")
            run.font.size = Pt(16)
            font = run.font
            font.name = '黑体'
            font._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')

            contents = list(set(iv))
            paragraph_content = file.add_paragraph()
            paragraph_content_ = file.add_paragraph()
            run_ = paragraph_content.add_run("，".join(contents))
            font = run_.font
            font.size = Pt(14)
            font.name = '华文宋体'
            font._element.rPr.rFonts.set(qn('w:eastAsia'), u'华文宋体')

            run_2 = paragraph_content_.add_run('\n平顶山市公安局卫东分局\n二〇二三年七月二十七日')
            paragraph_content_.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run_2.font.size = Pt(14)
            font = run_2.font
            font.name = '华文宋体'
            font._element.rPr.rFonts.set(qn('w:eastAsia'), u'华文宋体')

            # 保存
            bankcard_sum = sum([int(file.split('-')[2].split('张')[0]) for file in os.listdir(f'./{folder_num}') if '文书附件' in file])
            if len(list(set(iv))) < 100 and len(list(set(iv))) > 90:
                os.mkdir(f'./{folder_num + 1}')
                file.save(f"./{folder_num + 1}/文书附件-{ik}-{len(list(set(iv)))}张.docx")
                folder_num += 1
                # bankcard_sum = sum([int(re.findall(r'\d+', file)[0]) for file in os.listdir(f'./{folder_num}') if '文书附件' in file])
            else:
                no100_df.append({'content': file, 'filename': '文书附件-' + str(ik) + '-' + str(len(list(set(iv)))) + '张.docx',
                                 'len': len(list(set(iv)))})

    no100_df.sort(key=itemgetter('len'), reverse=True)

    os.mkdir(f'./{folder_num + 1}')
    # no100_df[0]['content'].save(f"./{folder_num + 1}/{no100_df[0]['filename']}")
    no100_left = no100_df[: int(len(no100_df) / 2)]
    no100_right = no100_df[int(len(no100_df) / 2): ]
    no100_right.sort(key=itemgetter('len'), reverse=False)

    fail_save = []
    if len(no100_left) == len(no100_right):
        for i in range(len(no100_right)):
            bankcard_sum = sum([int(file.split('-')[2].split('张')[0]) for file in os.listdir(f'./{folder_num + 1}') if '文书附件' in file])
            if no100_left[i]['len'] <= 100 - bankcard_sum:
                no100_left[i]['content'].save(f"./{folder_num + 1}/{no100_left[i]['filename']}")
            else:
                fail_save.append(no100_left[i])
            if no100_right[i]['len'] <= 100 - bankcard_sum:
                no100_right[i]['content'].save(f"./{folder_num + 1}/{no100_right[i]['filename']}")
            else:
                fail_save.append(no100_right[i])
        if len(fail_save) != 0:
            for temp in fail_save:
                bankcard_sum = sum([int(file.split('-')[2].split('张')[0]) for file in os.listdir(f'./{folder_num + 1}') if '文书附件' in file])
                if temp['len'] <= 100 - bankcard_sum:
                    temp['content'].save(f"./{folder_num + 1}/{temp['filename']}")
                else:
                    folder_num += 1
                    os.mkdir(f'./{folder_num + 1}')
                    temp['content'].save(f"./{folder_num + 1}/{temp['filename']}")
    else:
        for i in range(len(no100_left)):
            bankcard_sum = sum([int(file.split('-')[2].split('张')[0]) for file in os.listdir(f'./{folder_num + 1}') if '文书附件' in file])
            if no100_left[i]['len'] <= 100 - bankcard_sum:
                no100_left[i]['content'].save(f"./{folder_num + 1}/{no100_left[i]['filename']}")
            else:
                fail_save.append(no100_left[i])
            if no100_right[i]['len'] <= 100 - bankcard_sum:
                no100_right[i]['content'].save(f"./{folder_num + 1}/{no100_right[i]['filename']}")
            else:
                fail_save.append(no100_right[i])

        if len(fail_save) != 0:
            for temp in fail_save:
                bankcard_sum = sum([int(file.split('-')[2].split('张')[0]) for file in os.listdir(f'./{folder_num + 1}') if '文书附件' in file])
                if temp['len'] <= 100 - bankcard_sum:
                    temp['content'].save(f"./{folder_num + 1}/{temp['filename']}")
                else:
                    folder_num += 1
                    os.mkdir(f'./{folder_num + 1}')
                    temp['content'].save(f"./{folder_num + 1}/{temp['filename']}")
        bankcard_sum = sum([int(file.split('-')[2].split('张')[0]) for file in os.listdir(f'./{folder_num + 1}') if '文书附件' in file])
        if no100_right[-1]['len'] <= 100 - bankcard_sum:
            no100_right[-1]['content'].save(f"./{folder_num + 1}/{no100_right[-1]['filename']}")
        else:
            folder_num += 1
            os.mkdir(f'./{folder_num + 1}')
            no100_right[-1]['content'].save(f"./{folder_num + 1}/{no100_right[-1]['filename']}")


    folders = [file for file in os.listdir('./') if '.' not in file]
    for folder in folders:
        data = []
        df_acc = []
        no = 1
        bankcard_sum = sum([int(file.split('-')[2].split('张')[0]) for file in os.listdir(f'./{folder}') if '文书附件' in file])
        for d in os.listdir(f'./{folder}'):
            if '文书附件' in d:
                doc = docx.Document(f'./{folder}/{d}')
                obj_figure = re.compile(r'\d+', re.S)
                accounts = [acc.text for acc in doc.paragraphs][1].split('，')
                bankName = [acc.text for acc in doc.paragraphs][0].split('：')[1].split('-')[0]
                try:
                    bankName = bankName.replace(obj_figure.search(bankName).group(), '')
                except:
                    pass
                for acc in accounts:
                    temp = df.loc[df['被查账/卡号'] == acc]
                    df_acc += [temp]

                if len(accounts) <= 10:
                    accountsStr = ','.join(accounts)
                    data.append({'no': no, 'bankName': bankName, 'cardNo': accountsStr})
                    no += 1
                else:
                    data.append({'no': no, 'bankName': bankName,
                                 'cardNo': f"{accounts[0]}等{len(accounts)}个涉案账户，详见相关账户表"})
                    no += 1
        # data.append({'bankNum': no, 'accountNum': bankcard_sum})
        pd.concat(df_acc).to_excel(f'./{folder}/待调单-{folder}.xls', index=False, engine='openpyxl')
        dataApply = {'items': data, 'bankNum': no - 1, 'accountNum': bankcard_sum}
        tpl = DocxTemplate('./采取查询手段申请表_模板.docx')
        for idx, d in enumerate([dataApply]):
            context = d
            tpl.render(context)
            tpl.save(f'./{folder}/采取查询手段申请表{folder}.docx')