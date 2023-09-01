# -*- coding: utf-8 -*-
# @Author  :  connor
# @Time    :  2023/8/4 16:44
# @File    :  word2excel.py
# @IDE     :  PyCharm

"""
word2excel
"""
import docx
import xlwt


def write_excel(path):


    doc = docx.Document(docFile)
    workbook = xlwt.Workbook(encoding='utf-8')
    # 创建sheet
    data_sheet = workbook.add_sheet('sheet1')
    index = 1

    for i in range(0, len(doc.paragraphs), 13):
        if i + 12 > len(doc.paragraphs):
            break
        no = doc.paragraphs[i].text
        statu = doc.paragraphs[i + 1].text
        name = doc.paragraphs[i + 2].text
        sex = doc.paragraphs[i + 3].text
        cardNo = doc.paragraphs[i + 4].text
        personalAddress = doc.paragraphs[i + 5].text
        caseType = doc.paragraphs[i + 6].text
        detail = doc.paragraphs[i + 7].text
        currentAddress = doc.paragraphs[i + 8].text
        date = doc.paragraphs[i + 9].text
        cardNoAddress = doc.paragraphs[i + 10].text
        detailDate = doc.paragraphs[i + 11].text
        depart = doc.paragraphs[i + 12].text


        data_sheet.write(index, 0, index, xlwt.XFStyle())
        data_sheet.write(index, 1, no, xlwt.XFStyle())
        data_sheet.write(index, 2, statu, xlwt.XFStyle())
        data_sheet.write(index, 3, name, xlwt.XFStyle())
        data_sheet.write(index, 4, sex, xlwt.XFStyle())
        data_sheet.write(index, 5, cardNo, xlwt.XFStyle())
        data_sheet.write(index, 6, personalAddress, xlwt.XFStyle())
        data_sheet.write(index, 7, caseType, xlwt.XFStyle())
        data_sheet.write(index, 8, detail, xlwt.XFStyle())
        data_sheet.write(index, 9, currentAddress, xlwt.XFStyle())
        data_sheet.write(index, 10, date, xlwt.XFStyle())
        data_sheet.write(index, 11, cardNoAddress, xlwt.XFStyle())
        data_sheet.write(index, 12, detailDate, xlwt.XFStyle())
        data_sheet.write(index, 13, depart, xlwt.XFStyle())
        index += 1
        print("--------------------------")

    workbook.save(path)

if __name__ == '__main__':

    docFile = "./xx办理XX人员.docx"
    # 设置路径
    path = './xx办理XX人员.xls'  # 写入路径
    write_excel(path)
    print(u'创建文件成功')

